import os
import sys
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from PIL import Image
import tempfile
from tqdm import tqdm

tesseract_cmd = '/usr/bin/tesseract'

if not os.path.exists(tesseract_cmd):
    print(f"Tesseract not found at {tesseract_cmd}. Please install Tesseract OCR and update the path in the script.")
    sys.exit(1)

pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

def process_page(page, page_num, doc):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
        temp_image_path = temp_file.name
    
    page.save(temp_image_path, 'PNG')
    
    try:
        text = pytesseract.image_to_string(Image.open(temp_image_path), lang='mar')
    except pytesseract.TesseractError as e:
        print(f"\nError during OCR on page {page_num}: {e}")
        text = f"OCR failed for page {page_num}."
    
    doc.add_paragraph(text)
    doc.add_page_break()
    
    os.unlink(temp_image_path)

def pdf_to_word_devanagari(pdf_path, output_folder, dpi=150, batch_size=10):
    os.makedirs(output_folder, exist_ok=True)
    doc = Document()

    try:
        print("Counting total pages...")
        total_pages = len(convert_from_path(pdf_path, dpi=dpi))
        print(f"Total pages: {total_pages}")

        for batch_start in range(1, total_pages + 1, batch_size):
            batch_end = min(batch_start + batch_size - 1, total_pages)
            pages = convert_from_path(pdf_path, dpi=dpi, first_page=batch_start, last_page=batch_end)
            
            for i, page in enumerate(pages):
                page_num = batch_start + i
                process_page(page, page_num, doc)

    except Exception as e:
        print(f"\nError processing PDF: {e}")
        raise

    if len(doc.paragraphs) > 0:
        doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

    output_filename = "extracted_text.docx"
    doc_path = os.path.join(output_folder, output_filename)
    doc.save(doc_path)
    print(f"\nWord document saved at: {doc_path}")
    return output_filename